from pathlib import Path
from typing import Optional
import PyPDF2
import pdfplumber
import os
from ..core.models import Document
from ..core.utils import get_file_size, compute_text_hash, clean_text
from ..core.logging_config import get_logger
from ..core.exceptions import DocumentIngestionError

logger = get_logger(__name__)


class DocumentLoader:
    """Base class for document loaders."""
    
    @staticmethod
    def load_txt(file_path: str) -> Document:
        """Load a text file.
        
        Args:
            file_path: Path to the text file
            
        Returns:
            Document object
        """
        try:
            path = Path(file_path)
            
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Clean content
            content = clean_text(content)
            
            # Generate metadata
            content_hash = compute_text_hash(content)
            doc_id = Document.generate_doc_id(str(path), content_hash)
            
            doc = Document(
                doc_id=doc_id,
                content=content,
                source=str(path.absolute()),
                filename=path.name,
                file_type='txt',
                file_size=get_file_size(file_path),
                content_hash=content_hash,
                metadata={
                    'loader': 'txt'
                }
            )
            
            logger.info(f"Loaded TXT file: {path.name}")
            return doc
            
        except Exception as e:
            raise DocumentIngestionError(f"Failed to load TXT file {file_path}: {str(e)}")
    
    @staticmethod
    def load_pdf(file_path: str, use_pdfplumber: bool = True) -> Document:
        """Load a PDF file.
        
        Args:
            file_path: Path to the PDF file
            use_pdfplumber: Use pdfplumber (better text extraction) vs PyPDF2
            
        Returns:
            Document object
        """
        try:
            path = Path(file_path)
            
            if use_pdfplumber:
                content = DocumentLoader._extract_pdf_pdfplumber(path)
            else:
                content = DocumentLoader._extract_pdf_pypdf2(path)
            
            # Clean content
            content = clean_text(content)
            
            # Generate metadata
            content_hash = compute_text_hash(content)
            doc_id = Document.generate_doc_id(str(path), content_hash)
            
            doc = Document(
                doc_id=doc_id,
                content=content,
                source=str(path.absolute()),
                filename=path.name,
                file_type='pdf',
                file_size=get_file_size(file_path),
                content_hash=content_hash,
                metadata={
                    'loader': 'pdfplumber' if use_pdfplumber else 'pypdf2'
                }
            )
            
            logger.info(f"Loaded PDF file: {path.name}")
            return doc
            
        except Exception as e:
            raise DocumentIngestionError(f"Failed to load PDF file {file_path}: {str(e)}")
    
    @staticmethod
    def _extract_pdf_pypdf2(path: Path) -> str:
        """Extract text from PDF using PyPDF2."""
        text_parts = []
        
        with open(path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        
        return '\n\n'.join(text_parts)
    
    @staticmethod
    def _extract_pdf_pdfplumber(path: Path) -> str:
        """Extract text from PDF using pdfplumber (better quality)."""
        text_parts = []
        
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        
        return '\n\n'.join(text_parts)
    
    @staticmethod
    def load_markdown(file_path: str) -> Document:
        """Load a Markdown file.
        
        Args:
            file_path: Path to the Markdown file
            
        Returns:
            Document object
        """
        try:
            path = Path(file_path)
            
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Clean content (preserve markdown formatting)
            content = clean_text(content)
            
            # Generate metadata
            content_hash = compute_text_hash(content)
            doc_id = Document.generate_doc_id(str(path), content_hash)
            
            doc = Document(
                doc_id=doc_id,
                content=content,
                source=str(path.absolute()),
                filename=path.name,
                file_type='md',
                file_size=get_file_size(file_path),
                content_hash=content_hash,
                metadata={
                    'loader': 'markdown'
                }
            )
            
            logger.info(f"Loaded Markdown file: {path.name}")
            return doc
            
        except Exception as e:
            raise DocumentIngestionError(f"Failed to load Markdown file {file_path}: {str(e)}")
    
    @staticmethod
    def load_docx(file_path: str) -> Document:
        """Load a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Document object
        """
        try:
            from docx import Document as DocxDocument
            
            path = Path(file_path)
            doc = DocxDocument(str(path))
            
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = '\t'.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            
            content = '\n\n'.join(paragraphs)
            content = clean_text(content)
            
            content_hash = compute_text_hash(content)
            doc_id = Document.generate_doc_id(str(path), content_hash)
            
            return Document(
                doc_id=doc_id,
                content=content,
                source=str(path.absolute()),
                filename=path.name,
                file_type='docx',
                file_size=get_file_size(file_path),
                content_hash=content_hash,
                metadata={'loader': 'python-docx'}
            )
        except ImportError:
            raise DocumentIngestionError("python-docx is required to load DOCX files. Install with: pip install python-docx")
        except Exception as e:
            raise DocumentIngestionError(f"Failed to load DOCX file {file_path}: {str(e)}")
    
    @staticmethod
    def load_excel(file_path: str) -> Document:
        """Load an Excel file (xlsx, xls, csv).
        
        Args:
            file_path: Path to the Excel file
            
        Returns:
            Document object
        """
        try:
            import pandas as pd
            
            path = Path(file_path)
            ext = path.suffix.lower()
            
            if ext == '.csv':
                df = pd.read_csv(str(path))
                sheets_text = [f"CSV Data ({len(df)} rows, {len(df.columns)} columns):\n\nColumns: {', '.join(df.columns.tolist())}\n\n{df.to_string(index=False, max_rows=500)}"]
            else:
                xls = pd.ExcelFile(str(path))
                sheets_text = []
                for sheet_name in xls.sheet_names:
                    df = pd.read_excel(xls, sheet_name=sheet_name)
                    sheet_info = f"Sheet: {sheet_name} ({len(df)} rows, {len(df.columns)} columns)\nColumns: {', '.join(df.columns.astype(str).tolist())}\n\n{df.to_string(index=False, max_rows=500)}"
                    sheets_text.append(sheet_info)
            
            content = '\n\n---\n\n'.join(sheets_text)
            content = clean_text(content)
            
            content_hash = compute_text_hash(content)
            doc_id = Document.generate_doc_id(str(path), content_hash)
            
            return Document(
                doc_id=doc_id,
                content=content,
                source=str(path.absolute()),
                filename=path.name,
                file_type=ext.lstrip('.'),
                file_size=get_file_size(file_path),
                content_hash=content_hash,
                metadata={'loader': 'pandas'}
            )
        except Exception as e:
            raise DocumentIngestionError(f"Failed to load Excel file {file_path}: {str(e)}")
    
    @staticmethod
    def load_image(file_path: str) -> Document:
        """Load an image file and extract any text via OCR or create a description.
        
        Args:
            file_path: Path to the image file
            
        Returns:
            Document object
        """
        try:
            from PIL import Image
            
            path = Path(file_path)
            img = Image.open(str(path))
            
            # Build metadata description
            width, height = img.size
            mode = img.mode
            img_format = img.format or path.suffix.lstrip('.').upper()
            
            content_parts = [
                f"Image file: {path.name}",
                f"Format: {img_format}",
                f"Dimensions: {width}x{height} pixels",
                f"Color mode: {mode}",
            ]
            
            # Try OCR with pytesseract if available
            try:
                import pytesseract
                ocr_text = pytesseract.image_to_string(img).strip()
                if ocr_text:
                    content_parts.append(f"\nExtracted text (OCR):\n{ocr_text}")
            except ImportError:
                content_parts.append("\nNote: Install pytesseract for OCR text extraction from images.")
            except Exception as ocr_err:
                logger.warning(f"OCR failed for {path.name}: {ocr_err}")
            
            # EXIF data
            exif = img.getexif()
            if exif:
                exif_items = []
                for tag_id, value in list(exif.items())[:20]:
                    try:
                        from PIL.ExifTags import TAGS
                        tag_name = TAGS.get(tag_id, tag_id)
                        exif_items.append(f"{tag_name}: {value}")
                    except Exception:
                        pass
                if exif_items:
                    content_parts.append(f"\nEXIF metadata:\n" + '\n'.join(exif_items))
            
            content = '\n'.join(content_parts)
            content_hash = compute_text_hash(content)
            doc_id = Document.generate_doc_id(str(path), content_hash)
            
            return Document(
                doc_id=doc_id,
                content=content,
                source=str(path.absolute()),
                filename=path.name,
                file_type=path.suffix.lstrip('.').lower(),
                file_size=get_file_size(file_path),
                content_hash=content_hash,
                metadata={
                    'loader': 'pillow',
                    'width': width,
                    'height': height,
                    'format': img_format,
                    'is_image': True
                }
            )
        except ImportError:
            raise DocumentIngestionError("Pillow is required to load image files. Install with: pip install Pillow")
        except Exception as e:
            raise DocumentIngestionError(f"Failed to load image file {file_path}: {str(e)}")
    
    @staticmethod
    def load_audio(file_path: str) -> Document:
        """Load an audio file and extract metadata/transcription.
        
        Args:
            file_path: Path to the audio file
            
        Returns:
            Document object
        """
        try:
            path = Path(file_path)
            file_size = get_file_size(file_path)
            ext = path.suffix.lower()
            
            content_parts = [
                f"Audio file: {path.name}",
                f"Format: {ext.lstrip('.')}",
                f"File size: {file_size / (1024*1024):.2f} MB",
            ]
            
            # Try to get audio duration/metadata with mutagen
            try:
                import mutagen
                audio = mutagen.File(str(path))
                if audio and audio.info:
                    duration_secs = getattr(audio.info, 'length', 0)
                    if duration_secs:
                        mins, secs = divmod(int(duration_secs), 60)
                        content_parts.append(f"Duration: {mins}m {secs}s")
                    sample_rate = getattr(audio.info, 'sample_rate', None)
                    if sample_rate:
                        content_parts.append(f"Sample rate: {sample_rate} Hz")
                    channels = getattr(audio.info, 'channels', None)
                    if channels:
                        content_parts.append(f"Channels: {channels}")
            except ImportError:
                pass
            except Exception as meta_err:
                logger.warning(f"Audio metadata extraction failed: {meta_err}")
            
            # Try speech-to-text with SpeechRecognition for WAV files
            try:
                import speech_recognition as sr
                if ext == '.wav':
                    recognizer = sr.Recognizer()
                    with sr.AudioFile(str(path)) as source:
                        audio_data = recognizer.record(source, duration=300)  # max 5 mins
                    transcript = recognizer.recognize_google(audio_data)
                    if transcript:
                        content_parts.append(f"\nTranscription:\n{transcript}")
            except ImportError:
                content_parts.append("\nNote: Install SpeechRecognition for audio transcription.")
            except Exception as stt_err:
                logger.warning(f"Speech-to-text failed for {path.name}: {stt_err}")
                content_parts.append(f"\nNote: Automatic transcription was not available for this file.")
            
            content = '\n'.join(content_parts)
            content_hash = compute_text_hash(content)
            doc_id = Document.generate_doc_id(str(path), content_hash)
            
            return Document(
                doc_id=doc_id,
                content=content,
                source=str(path.absolute()),
                filename=path.name,
                file_type=ext.lstrip('.'),
                file_size=file_size,
                content_hash=content_hash,
                metadata={
                    'loader': 'audio',
                    'is_audio': True
                }
            )
        except DocumentIngestionError:
            raise
        except Exception as e:
            raise DocumentIngestionError(f"Failed to load audio file {file_path}: {str(e)}")
    
    @staticmethod
    def load_file(file_path: str) -> Document:
        """Load a file based on its extension.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Document object
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        
        if extension == '.txt':
            return DocumentLoader.load_txt(file_path)
        elif extension == '.pdf':
            return DocumentLoader.load_pdf(file_path)
        elif extension in ['.md', '.markdown']:
            return DocumentLoader.load_markdown(file_path)
        elif extension in ['.docx', '.doc']:
            return DocumentLoader.load_docx(file_path)
        elif extension in ['.xlsx', '.xls', '.csv']:
            return DocumentLoader.load_excel(file_path)
        elif extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp', '.tiff', '.svg']:
            return DocumentLoader.load_image(file_path)
        elif extension in ['.mp3', '.wav', '.ogg', '.flac', '.m4a', '.aac', '.wma']:
            return DocumentLoader.load_audio(file_path)
        else:
            raise DocumentIngestionError(f"Unsupported file type: {extension}")
